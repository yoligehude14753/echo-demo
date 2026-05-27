"""Office 产物的 Python 兜底生成器（移植自 meetly，加入流程图 PNG 嵌入）。

Claude Code subprocess 在以下场景会失败：
- 本机没装 ``claude`` binary
- proxy 不可达
- 远程 LLM 长尾抖动

「产物不存在」=「会议结束没纪要」是业务上的硬故障；用 python-docx 直接生成最低合规版
（标题 + 章节 + 表格 + 流程图 PNG），保证总有产物落地。
"""
from __future__ import annotations

import html
from pathlib import Path
from typing import Any


def fallback_docx(
    target_path: Path,
    *,
    title: str,
    abstract: str = "",
    summary_md: str = "",
    decisions: list[dict[str, Any]] | None = None,
    todos: list[dict[str, Any]] | None = None,
    topics: list[dict[str, Any]] | None = None,
    flow_png_path: Path | None = None,
    participants: list[str] | None = None,
    time_range: str = "",
) -> None:
    """用 python-docx 生成保底 Word 文档。

    与 meetly 版本相比新增：
    - abstract 单独一段（不与 summary_md 重复）
    - topics 单独章节
    - flow_png_path 嵌入到「会议流程」章节
    - 元信息（时间/参会人）放在标题下
    """
    from docx import Document
    from docx.shared import Inches, Pt

    decisions = decisions or []
    todos = todos or []
    topics = topics or []
    participants = participants or []

    d = Document()
    style = d.styles["Normal"]
    style.font.size = Pt(11)

    # 标题 + 元信息
    d.add_heading(title or "会议纪要", level=0)
    if time_range or participants:
        meta = d.add_paragraph()
        if time_range:
            meta.add_run(f"时间：{time_range}\n").italic = True
        if participants:
            meta.add_run(f"参会人：{'、'.join(participants)}").italic = True

    # 摘要
    if abstract:
        d.add_heading("摘要", level=1)
        d.add_paragraph(abstract)

    # 会议流程（含 PNG）
    if flow_png_path and flow_png_path.exists():
        d.add_heading("会议流程", level=1)
        try:
            d.add_picture(str(flow_png_path), width=Inches(6))
        except Exception:
            d.add_paragraph("（流程图渲染失败，请查看 HTML 预览）")
    elif flow_png_path is not None:
        d.add_heading("会议流程", level=1)
        d.add_paragraph("（流程图未渲染：mmdc 不可用或返回异常，请查看 HTML 预览）")

    # 决议
    if decisions:
        d.add_heading("会议决议", level=1)
        for i, dec in enumerate(decisions, 1):
            p = d.add_paragraph(style="List Number")
            statement = dec.get("statement") or dec.get("decision") or ""
            run = p.add_run(str(statement))
            run.bold = True
            ra = dec.get("rationale")
            if ra:
                d.add_paragraph(f"  依据：{ra}")
            ip = dec.get("impact")
            if ip:
                d.add_paragraph(f"  影响：{ip}")

    # 待办（真实 docx table）
    if todos:
        d.add_heading("待办事项", level=1)
        table = d.add_table(rows=1, cols=4)
        try:
            table.style = "Light List Accent 1"
        except KeyError:
            pass
        hdr = table.rows[0].cells
        hdr[0].text = "任务"
        hdr[1].text = "负责人"
        hdr[2].text = "截止"
        hdr[3].text = "优先级"
        priority_label = {"high": "高", "med": "中", "low": "低"}
        for t in todos:
            row = table.add_row().cells
            row[0].text = str(t.get("task") or "")
            row[1].text = str(t.get("owner") or "未指派")
            row[2].text = str(t.get("due") or "TBD")
            row[3].text = priority_label.get(str(t.get("priority") or "med"), "中")

    # 话题
    if topics:
        d.add_heading("话题展开", level=1)
        for tp in topics:
            name = str(tp.get("name") or "").strip()
            time_range_str = str(tp.get("time_range") or "").strip()
            if name:
                heading_text = f"{name}（{time_range_str}）" if time_range_str else name
                d.add_heading(heading_text, level=2)
            key_points = tp.get("key_points") or []
            for kp in key_points:
                kp_str = str(kp).strip()
                if kp_str:
                    d.add_paragraph(kp_str, style="List Bullet")

    # 完整纪要 markdown 渲染（简化：仅支持 # / ## / - 三种标记）
    if summary_md:
        d.add_heading("完整纪要", level=1)
        for line in summary_md.split("\n"):
            line = line.rstrip()
            stripped = line.strip()
            if stripped.startswith("### "):
                d.add_heading(stripped[4:].strip(), level=3)
            elif stripped.startswith("## "):
                d.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith("# "):
                d.add_heading(stripped[2:].strip(), level=2)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                d.add_paragraph(stripped[2:].strip(), style="List Bullet")
            elif stripped:
                d.add_paragraph(stripped)

    target_path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(target_path))


def fallback_html_minimal(
    target_path: Path,
    *,
    title: str,
    abstract: str,
    summary_md: str,
    decisions: list[dict[str, Any]],
    todos: list[dict[str, Any]],
) -> None:
    """极简 HTML 兜底——只在 Jinja2 模板渲染崩溃时使用，正常路径走 renderers/html.py。"""

    def esc(s: Any) -> str:
        return html.escape(str(s or ""))

    rows_todo = "".join(
        f"<tr><td>{esc(t.get('task'))}</td><td>{esc(t.get('owner') or '未指派')}</td>"
        f"<td>{esc(t.get('due') or 'TBD')}</td>"
        f"<td>{esc(t.get('priority') or 'med')}</td></tr>"
        for t in todos
    )
    rows_dec = "".join(
        f"<li><strong>{esc(d.get('statement') or d.get('decision'))}</strong></li>"
        for d in decisions
    )

    out = f"""<!doctype html>
<html lang="zh"><head><meta charset="utf-8"><title>{esc(title)}</title>
<style>body{{font-family:sans-serif;max-width:920px;margin:40px auto;padding:0 24px}}
table{{width:100%;border-collapse:collapse}}th,td{{padding:8px;border-bottom:1px solid #eee}}</style>
</head><body>
<h1>{esc(title or '会议纪要')}</h1>
<p>{esc(abstract)}</p>
<h2>决议</h2><ol>{rows_dec or '<li>（无）</li>'}</ol>
<h2>待办</h2><table><thead><tr><th>任务</th><th>负责人</th><th>截止</th><th>优先级</th></tr></thead>
<tbody>{rows_todo or '<tr><td colspan=4>（无）</td></tr>'}</tbody></table>
<h2>完整纪要</h2><pre style="white-space:pre-wrap">{esc(summary_md)}</pre>
</body></html>"""
    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_path.write_text(out, encoding="utf-8")
