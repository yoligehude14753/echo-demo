"""真链路验证：不同体裁的 Word/Excel 生成后，结构是否真的随内容自适应。

对每个 (类型, brief) 调真实 /artifacts/generate，再用 python-docx / openpyxl
读回产物，打印它的实际结构（docx: 标题层级 + 表格数；xlsx: sheet 名 + 每表列头 +
是否含公式），用来肉眼确认"章节/表结构按内容变了，且没退化"。

用法：.venv/bin/python scripts/stress/inspect_doc_structure.py
"""

from __future__ import annotations

import sys

import httpx

BASE = "http://127.0.0.1:8769"

CASES: list[tuple[str, str]] = [
    ("word", "整理今天的产品评审会议纪要：参会张三李四王五，讨论了登录改版和支付bug两个议题，定了下周一上线，待办：张三改UI、李四修支付"),
    ("word", "写一份公司端午节放假通知，放假6月7到9号共三天，值班电话保留，注意安全"),
    ("xlsx", "做一个新产品上线排期表，从需求评审到灰度发布分5个阶段，列出阶段、起止日期、负责人、状态"),
    ("xlsx", "做一个团建活动预算表，包含场地、餐饮、交通、礼品四类，每类有单价数量金额，最后要合计"),
]


def gen(client: httpx.Client, atype: str, brief: str) -> dict:
    r = client.post(
        f"{BASE}/artifacts/generate",
        json={"artifact_type": atype, "brief": brief},
        timeout=240.0,
    )
    r.raise_for_status()
    return r.json()


def inspect_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    headings = [
        f"{p.style.name}:{p.text[:24]}"
        for p in doc.paragraphs
        if p.style and p.style.name.startswith("Heading") and p.text.strip()
    ]
    return f"标题数={len(headings)} 表格数={len(doc.tables)} | 标题样例={headings[:6]}"


def inspect_xlsx(path: str) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path)
    parts = []
    has_formula = False
    for ws in wb.worksheets:
        headers = [c.value for c in next(ws.iter_rows(max_row=1), [])]
        for row in ws.iter_rows():
            for c in row:
                if isinstance(c.value, str) and c.value.startswith("="):
                    has_formula = True
        parts.append(f"[{ws.title}] 列={headers}")
    return f"sheets={wb.sheetnames} 含公式={has_formula} | {' ; '.join(parts)}"


def main() -> int:
    with httpx.Client() as client:
        for atype, brief in CASES:
            try:
                art = gen(client, atype, brief)
            except Exception as e:
                print(f"[{atype}] FAIL: {e}")
                continue
            path = art.get("file_path", "")
            variant = art.get("metadata", {}).get("skill_variant")
            try:
                summary = inspect_docx(path) if atype == "word" else inspect_xlsx(path)
            except Exception as e:
                summary = f"(inspect failed: {e})"
            print(f"\n[{atype}] variant={variant} size={art.get('size_bytes')}")
            print(f"  brief: {brief[:34]}…")
            print(f"  结构: {summary}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
