"""参考 docx 样式抽离单测（离线，自造参考件 → 抽离 → 断言）。"""

from __future__ import annotations

from pathlib import Path

import pytest
from app.use_cases.docx_style import extract_docx_style, extract_style_instructions
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


def _set_ea(run: object, font: str) -> None:
    run.font.name = font  # type: ignore[attr-defined]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]


def _make_reference(path: Path, *, red_head: bool) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin, sec.bottom_margin = Mm(37), Mm(35)
    sec.left_margin, sec.right_margin = Mm(28), Mm(26)

    if red_head:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("某某市人民政府")
        r.font.size = Pt(22)
        r.bold = True
        r.font.color.rgb = RGBColor(0xE0, 0x00, 0x00)
        _set_ea(r, "方正小标宋简体")

    # 大标题（居中，二号）
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run("关于测试样式抽离的通知")
    rt.font.size = Pt(22)
    rt.bold = True
    _set_ea(rt, "方正小标宋简体")

    # 一级标题（Heading 1，黑体三号）
    h1 = doc.add_heading(level=1)
    rh = h1.add_run("一、背景")
    rh.font.size = Pt(16)
    _set_ea(rh, "黑体")

    # 正文（仿宋三号，固定行距 28 磅）
    for _ in range(3):
        bp = doc.add_paragraph()
        bp.paragraph_format.line_spacing = Pt(28)
        rb = bp.add_run("这是正文段落，用于抽离正文字体与字号。")
        rb.font.size = Pt(16)
        _set_ea(rb, "仿宋_GB2312")

    doc.save(str(path))


@pytest.mark.unit
def test_extract_margins_and_fonts(tmp_path: Path) -> None:
    ref = tmp_path / "ref.docx"
    _make_reference(ref, red_head=False)
    spec = extract_docx_style(str(ref))

    assert round(spec.margins_mm["top"]) == 37
    assert round(spec.margins_mm["bottom"]) == 35
    assert round(spec.margins_mm["left"]) == 28
    assert round(spec.margins_mm["right"]) == 26

    assert spec.title is not None
    assert spec.title.east_asia == "方正小标宋简体"
    assert spec.title.size_pt == 22

    assert spec.body.east_asia == "仿宋_GB2312"
    assert spec.body.size_pt == 16
    assert spec.line_spacing_pt == 28

    assert 1 in spec.headings
    assert spec.headings[1].east_asia == "黑体"
    assert spec.has_red_head is False


@pytest.mark.unit
def test_extract_detects_red_head(tmp_path: Path) -> None:
    ref = tmp_path / "ref_red.docx"
    _make_reference(ref, red_head=True)
    spec = extract_docx_style(str(ref))
    assert spec.has_red_head is True


@pytest.mark.unit
def test_instructions_contain_extracted_style(tmp_path: Path) -> None:
    ref = tmp_path / "ref2.docx"
    _make_reference(ref, red_head=True)
    text = extract_style_instructions(str(ref))
    assert "仿宋_GB2312" in text
    assert "16pt" in text
    assert "上37" in text and "左28" in text
    assert "红头" in text
