"""知识库样式模板发现单测：从 RAG 命中本地 .docx → 抽样式 → 注入指令。"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import pytest
from app.schemas.rag import RagChunk
from app.use_cases.style_template import (
    merge_extra_instructions,
    resolve_docx_style_template,
)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


class _FakeRag:
    def __init__(self, chunks: list[RagChunk]) -> None:
        self._chunks = chunks

    async def query(self, query: str, *, top_k: int = 5) -> list[RagChunk]:
        return self._chunks

    async def ingest_pdf(self, *a: Any, **k: Any) -> str:  # pragma: no cover
        raise NotImplementedError

    async def ingest_file(self, *a: Any, **k: Any) -> str:  # pragma: no cover
        raise NotImplementedError

    async def ingest_meeting(self, *a: Any, **k: Any) -> str:  # pragma: no cover
        raise NotImplementedError

    async def ingest_ambient_segment(self, *a: Any, **k: Any) -> str:  # pragma: no cover
        raise NotImplementedError

    async def delete(self, doc_id: str) -> None:  # pragma: no cover
        raise NotImplementedError

    async def find_by_source_path(self, source_path: str) -> str | None:  # pragma: no cover
        return None

    async def list_docs(self) -> list[dict[str, object]]:  # pragma: no cover
        return []


def _make_docx(path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin, sec.bottom_margin = Mm(37), Mm(35)
    sec.left_margin, sec.right_margin = Mm(28), Mm(26)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("关于参考样式的通知")
    r.font.size = Pt(22)
    r.bold = True
    r.font.name = "方正小标宋简体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "方正小标宋简体")
    bp = doc.add_paragraph()
    rb = bp.add_run("正文内容")
    rb.font.size = Pt(16)
    rb.font.name = "仿宋_GB2312"
    rb._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    doc.save(str(path))


def _chunk(source_path: str, score: float) -> RagChunk:
    return RagChunk(
        doc_id="d",
        doc_title="ref",
        chunk_id="d-c0",
        text="关于……的通知 正文",
        score=score,
        metadata={"kind": "docx", "source": "workspace", "source_path": source_path},
    )


@pytest.mark.asyncio
@pytest.mark.unit
async def test_resolves_template_from_knowledge_base(tmp_path: Path) -> None:
    ref = tmp_path / "公文模板.docx"
    _make_docx(ref)
    rag = _FakeRag([_chunk(str(ref), 3.0), _chunk(str(ref), 2.0)])

    tpl = await resolve_docx_style_template(rag, "写一份放假通知")
    assert tpl is not None
    assert tpl.title == "公文模板.docx"
    assert "仿宋_GB2312" in tpl.instructions
    assert "上37" in tpl.instructions


@pytest.mark.asyncio
@pytest.mark.unit
async def test_no_docx_in_kb_returns_none() -> None:
    rag = _FakeRag(
        [
            RagChunk(
                doc_id="p",
                doc_title="pdf",
                chunk_id="p-c0",
                text="x",
                score=5.0,
                metadata={"kind": "pdf", "source_path": "/tmp/a.pdf"},
            )
        ]
    )
    assert await resolve_docx_style_template(rag, "写通知") is None


@pytest.mark.asyncio
@pytest.mark.unit
async def test_missing_file_returns_none(tmp_path: Path) -> None:
    rag = _FakeRag([_chunk(str(tmp_path / "不存在.docx"), 9.0)])
    assert await resolve_docx_style_template(rag, "写通知") is None


@pytest.mark.unit
def test_merge_extra_instructions() -> None:
    assert merge_extra_instructions(None, "S") == "S"
    assert merge_extra_instructions("", "S") == "S"
    merged = merge_extra_instructions("用户要求", "样式指令")
    assert merged.startswith("用户要求")
    assert merged.endswith("样式指令")
