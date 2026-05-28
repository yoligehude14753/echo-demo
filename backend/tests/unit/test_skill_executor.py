"""Skill 执行器单测：mock LLM，验证 7 种产物的代码路径。"""

from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

import pytest
from app.adapters.skill import SkillError, SkillExecutor
from app.adapters.skill.llm_skill import _make_title, _strip_code_fence
from app.adapters.skill.node_executor import _is_safe_node
from app.adapters.skill.python_executor import _is_safe_python
from app.config import Settings
from app.schemas.artifact import SUPPORTED_KINDS, GeneratedArtifact, normalize_kind
from app.schemas.llm import ChatMessage, LLMResponse, LLMUsage


class FakeLLM:
    def __init__(self, content: str) -> None:
        self.content = content
        self.last_messages: list[ChatMessage] | None = None

    async def chat(self, messages: list[ChatMessage], **_: Any) -> LLMResponse:
        self.last_messages = list(messages)
        return LLMResponse(
            content=self.content,
            model="MiniMax-M2.7",
            finish_reason="stop",
            usage=LLMUsage(prompt_tokens=10, completion_tokens=20, total_tokens=30),
            latency_ms=12.0,
        )

    async def chat_stream(self, _messages: list[ChatMessage], **_: Any):  # type: ignore[no-untyped-def]
        raise NotImplementedError
        yield  # pragma: no cover


def _settings(tmp_path: Path, *, use_legacy_html_pptx: bool = True) -> Settings:
    """通用测试 Settings；默认 ``use_legacy_html_pptx=True``，保持本文件历史
    用例（直写 Tailwind HTML / pptxgenjs js）仍走旧流水线。

    新 phase4-doc-skills 用例（Kami one-pager / IB deck JSON）显式传
    ``use_legacy_html_pptx=False``，对应测试在
    ``tests/unit/test_skill_doc_skills.py`` 集中。
    """
    return Settings(
        storage_dir=tmp_path,
        skill_executor_build_dir=tmp_path / "skill_build",
        skill_executor_timeout_s=30,
        skill_executor_max_tokens=80_000,
        use_legacy_html_pptx=use_legacy_html_pptx,
    )


@pytest.mark.unit
def test_strip_code_fence_python() -> None:
    code = _strip_code_fence("```python\nprint(1)\n```")
    assert code == "print(1)"


@pytest.mark.unit
def test_strip_code_fence_html() -> None:
    code = _strip_code_fence("```html\n<!DOCTYPE html>\n<html></html>\n```")
    assert code.startswith("<!DOCTYPE html>")
    assert code.endswith("<html></html>")


@pytest.mark.unit
def test_strip_code_fence_no_fence() -> None:
    code = _strip_code_fence("from docx import Document\ndoc = Document()")
    assert "from docx" in code


@pytest.mark.unit
def test_strip_code_fence_skips_leading_prose() -> None:
    """M2.7 thinking 残留：LLM 输出前导自然语言，跳到第一行代码。"""
    raw = (
        "We need to use openpyxl with multiple sheets and formulas.\n"
        "Plan: 4 sheets, 25 formulas.\n"
        "from openpyxl import Workbook\n"
        "wb = Workbook()\n"
    )
    assert _strip_code_fence(raw).startswith("from openpyxl")


@pytest.mark.unit
def test_strip_code_fence_skips_leading_prose_node() -> None:
    raw = (
        "I'll build a pptx using pptxgenjs.\n"
        "const PptxGenJS = require('pptxgenjs');\n"
        "const pres = new PptxGenJS();\n"
    )
    assert _strip_code_fence(raw).startswith("const PptxGenJS")


@pytest.mark.unit
def test_is_safe_python_rejects_network() -> None:
    ok, reason = _is_safe_python("import socket\ns = socket.socket()")
    assert not ok
    assert "socket" in reason


@pytest.mark.unit
def test_is_safe_python_accepts_docx() -> None:
    ok, _ = _is_safe_python("from docx import Document\ndoc = Document()\ndoc.save('output.docx')")
    assert ok


@pytest.mark.asyncio
@pytest.mark.unit
async def test_unsupported_artifact_type_raises(tmp_path: Path) -> None:
    skill = SkillExecutor(_settings(tmp_path))
    llm = FakeLLM("...")
    # 7 类产物 (html/pptx/word/xlsx/markdown/pdf/txt) + 别名都已支持；
    # 这里用一个明确不在 SUPPORTED_KINDS 的值（如 csv）
    with pytest.raises(SkillError, match="unsupported"):
        await skill.generate(llm=llm, artifact_type="csv", brief="x")


@pytest.mark.asyncio
@pytest.mark.unit
async def test_html_generation_minimal(tmp_path: Path) -> None:
    skill = SkillExecutor(_settings(tmp_path))
    html_code = (
        "<!DOCTYPE html>\n<html>\n<head><script src='https://cdn.tailwindcss.com'></script></head>\n"
        + "<body class='bg-slate-900 text-white p-8'>\n"
        + '<h1 class="text-2xl">测试报告</h1>\n'
        + "<div>"
        + "x" * 2000
        + "</div>\n"
        + "</body></html>"
    )
    llm = FakeLLM(html_code)
    art = await skill.generate(llm=llm, artifact_type="html", brief="生成 demo HTML")
    assert art.artifact_type == "html"
    assert art.file_path.endswith(".html")
    assert Path(art.file_path).exists()
    assert art.size_bytes > 1500
    # 校验系统 prompt 是 HTML 的
    assert llm.last_messages is not None
    assert "Tailwind" in llm.last_messages[0].content


@pytest.mark.asyncio
@pytest.mark.unit
async def test_html_too_short_raises(tmp_path: Path) -> None:
    skill = SkillExecutor(_settings(tmp_path))
    llm = FakeLLM("<html>x</html>")
    with pytest.raises(SkillError, match="too short"):
        await skill.generate(llm=llm, artifact_type="html", brief="x")


@pytest.mark.asyncio
@pytest.mark.unit
async def test_html_no_doctype_raises(tmp_path: Path) -> None:
    skill = SkillExecutor(_settings(tmp_path))
    llm = FakeLLM("just plain text but long enough to pass length check " * 100)
    with pytest.raises(SkillError, match="DOCTYPE"):
        await skill.generate(llm=llm, artifact_type="html", brief="x")


@pytest.mark.asyncio
@pytest.mark.unit
async def test_word_generation_executes_python_docx(tmp_path: Path) -> None:
    """生成最小可运行的 docx 代码 → 真实执行 → 验证文件大小。"""
    pytest.importorskip("docx")
    skill = SkillExecutor(_settings(tmp_path))
    code = (
        "from docx import Document\n"
        "doc = Document()\n"
        "doc.add_heading('echo-demo skill test', level=1)\n"
        "doc.add_paragraph('一段中文内容，验证 python-docx 链路。')\n"
        "doc.save('output.docx')\n"
    )
    llm = FakeLLM(code)
    art = await skill.generate(llm=llm, artifact_type="word", brief="生成 demo word")
    assert art.artifact_type == "word"
    assert art.file_path.endswith(".docx")
    assert Path(art.file_path).stat().st_size > 1000


@pytest.mark.asyncio
@pytest.mark.unit
async def test_xlsx_generation_executes_openpyxl(tmp_path: Path) -> None:
    pytest.importorskip("openpyxl")
    skill = SkillExecutor(_settings(tmp_path))
    code = (
        "from openpyxl import Workbook\n"
        "wb = Workbook()\n"
        "ws = wb.active\n"
        "ws.title = '假设'\n"
        "ws['A1'] = '增长率'\n"
        "ws['B1'] = 0.15\n"
        "ws2 = wb.create_sheet('预测')\n"
        "ws2['A1'] = '=假设!B1*100'\n"
        "wb.save('output.xlsx')\n"
    )
    llm = FakeLLM(code)
    art = await skill.generate(llm=llm, artifact_type="xlsx", brief="生成 demo excel")
    assert art.artifact_type == "xlsx"
    assert art.file_path.endswith(".xlsx")
    assert Path(art.file_path).stat().st_size > 1000


@pytest.mark.asyncio
@pytest.mark.unit
async def test_python_with_forbidden_import_raises(tmp_path: Path) -> None:
    skill = SkillExecutor(_settings(tmp_path))
    code = "import socket\nsocket.socket()\n"
    llm = FakeLLM(code)
    with pytest.raises(SkillError, match="forbidden"):
        await skill.generate(llm=llm, artifact_type="word", brief="x")


# ── PR-12 新增：ArtifactKind 别名归一 + pptx 路径 ─────────────────────────


@pytest.mark.unit
def test_supported_kinds_covers_all_aliases() -> None:
    # 13 个别名 -> 7 个 canonical kinds
    expected = {
        "ppt",
        "pptx",
        "word",
        "docx",
        "xlsx",
        "excel",
        "html",
        "markdown",
        "md",
        "mdown",
        "pdf",
        "txt",
        "text",
    }
    assert expected == SUPPORTED_KINDS


@pytest.mark.unit
def test_normalize_kind_pptx_alias() -> None:
    assert normalize_kind("ppt") == "pptx"
    assert normalize_kind("PPTX") == "pptx"
    assert normalize_kind("excel") == "xlsx"
    assert normalize_kind("xlsx") == "xlsx"
    assert normalize_kind("word") == "word"
    assert normalize_kind("docx") == "word"
    assert normalize_kind("html") == "html"


@pytest.mark.unit
def test_normalize_kind_markdown_pdf_txt_aliases() -> None:
    """P4-M3：新增 markdown/pdf/txt 别名归一化。"""
    assert normalize_kind("markdown") == "markdown"
    assert normalize_kind("md") == "markdown"
    assert normalize_kind("MD") == "markdown"
    assert normalize_kind("mdown") == "markdown"
    assert normalize_kind("pdf") == "pdf"
    assert normalize_kind("PDF") == "pdf"
    assert normalize_kind("txt") == "txt"
    assert normalize_kind("text") == "txt"
    assert normalize_kind("TEXT") == "txt"


@pytest.mark.unit
def test_normalize_kind_invalid_returns_empty() -> None:
    # P4-M3 之前 "pdf" 是 invalid；现在 invalid 走 csv / rtf 等
    assert normalize_kind("csv") == ""
    assert normalize_kind("rtf") == ""
    assert normalize_kind("") == ""


@pytest.mark.unit
def test_make_title_short() -> None:
    assert _make_title("生成英伟达 2025 Q3 财报分析") == "生成英伟达 2025 Q3 财报分析"


@pytest.mark.unit
def test_make_title_long_truncates_with_ellipsis() -> None:
    brief = "英伟达 2025 财年第三季度业绩超预期" * 5
    title = _make_title(brief, max_len=40)
    assert len(title) <= 41  # 40 chars + 1 ellipsis
    assert title.endswith("…")


@pytest.mark.unit
def test_make_title_collapses_whitespace() -> None:
    assert _make_title("  生成  \n\t  HTML   报告  ") == "生成 HTML 报告"


@pytest.mark.unit
def test_make_title_empty() -> None:
    assert _make_title("") == ""
    assert _make_title("   \n\t  ") == ""


@pytest.mark.unit
def test_generated_artifact_title_default_empty() -> None:
    """旧 fixture 不传 title 时默认为空字符串，避免破坏下游测试。"""
    a = GeneratedArtifact(
        artifact_id="x",
        artifact_type="html",
        file_path="/tmp/x.html",
        mime_type="text/html",
        size_bytes=100,
        generation_latency_ms=1.0,
        model="m",
    )
    assert a.title == ""


@pytest.mark.unit
def test_is_safe_node_rejects_child_process() -> None:
    ok, reason = _is_safe_node("const cp = require('child_process'); cp.exec('ls');")
    assert not ok
    assert "child_process" in reason


@pytest.mark.unit
def test_is_safe_node_rejects_fs() -> None:
    ok, reason = _is_safe_node("const fs = require('fs'); fs.readFileSync('x');")
    assert not ok
    assert "fs" in reason


@pytest.mark.unit
def test_is_safe_node_accepts_pptxgenjs() -> None:
    ok, _ = _is_safe_node(
        "const PptxGenJS = require('pptxgenjs');\n"
        "const pres = new PptxGenJS();\n"
        "pres.addSlide().addText('hi');\n"
        "pres.writeFile({ fileName: 'output.pptx' });\n"
    )
    assert ok


@pytest.mark.asyncio
@pytest.mark.unit
async def test_pptx_node_missing_marks_failure(tmp_path: Path) -> None:
    """node 不存在时返回 SkillError，不抛运行时异常。"""
    skill = SkillExecutor(
        Settings(
            storage_dir=tmp_path,
            skill_executor_build_dir=tmp_path / "skill_build",
            skill_executor_timeout_s=10,
            skill_executor_max_tokens=80_000,
            skill_node_bin="/non/existent/node-binary-xyz",
            use_legacy_html_pptx=True,  # 测试 legacy pptxgenjs 路径下 node 缺失分支
        )
    )
    code = (
        "const PptxGenJS = require('pptxgenjs');\n"
        "const pres = new PptxGenJS();\n"
        "pres.addSlide().addText('hi');\n"
        "pres.writeFile({ fileName: 'output.pptx' });\n"
    )
    llm = FakeLLM(code)
    with pytest.raises(SkillError):
        await skill.generate(llm=llm, artifact_type="pptx", brief="x")


@pytest.mark.asyncio
@pytest.mark.unit
async def test_pptx_forbidden_token_raises(tmp_path: Path) -> None:
    skill = SkillExecutor(_settings(tmp_path))
    code = "const cp = require('child_process');\npres.writeFile({fileName:'x.pptx'});"
    llm = FakeLLM(code)
    with pytest.raises(SkillError, match=r"forbidden|execution failed"):
        await skill.generate(llm=llm, artifact_type="pptx", brief="x")


@pytest.mark.asyncio
@pytest.mark.integration
@pytest.mark.skipif(
    shutil.which("node") is None or shutil.which("npm") is None,
    reason="node / npm 未在 PATH",
)
@pytest.mark.skipif(
    "CI" in __import__("os").environ and "ECHO_RUN_NODE_INSTALL" not in __import__("os").environ,
    reason="CI 默认不跑 npm install pptxgenjs（耗时不稳定，已被 yunwu E2E 覆盖）",
)
async def test_pptx_generation_executes_pptxgenjs(tmp_path: Path) -> None:
    """mock LLM 返回最小可运行的 pptxgenjs 代码 → 真跑 node → 验证 .pptx 文件.

    注：此测试需要 ``npm install pptxgenjs``（首次约 80 MB），网络不稳时易卡 CI。
    单测层只保留 _is_safe_node / normalize_kind / node 缺失分支等纯 Python 路径。
    完整 pptxgenjs 链路由 ``tests/integration/test_skill_e2e_yunwu.py::test_skill_pptx_real_yunwu``
    在 integration 阶段覆盖（本地或夜间 runner 跑）。
    """
    skill = SkillExecutor(_settings(tmp_path))
    code = (
        "const PptxGenJS = require('pptxgenjs');\n"
        "const pres = new PptxGenJS();\n"
        "pres.layout = 'LAYOUT_WIDE';\n"
        "const s1 = pres.addSlide();\n"
        "s1.addText('echo demo 1', { x: 0.5, y: 0.5, fontSize: 28 });\n"
        "const s2 = pres.addSlide();\n"
        "s2.addText('echo demo 2', { x: 0.5, y: 0.5, fontSize: 28 });\n"
        "const s3 = pres.addSlide();\n"
        "s3.addText('echo demo 3', { x: 0.5, y: 0.5, fontSize: 28 });\n"
        "pres.writeFile({ fileName: 'output.pptx' });\n"
    )
    llm = FakeLLM(code)
    art = await skill.generate(llm=llm, artifact_type="pptx", brief="生成 demo pptx")
    assert art.artifact_type == "pptx"
    assert art.file_path.endswith(".pptx")
    assert Path(art.file_path).stat().st_size > 8_000
    assert int(art.metadata.get("slide_count_hint", "0")) >= 3


# ── P4-M3 新增：markdown / txt / pdf 三种产物 + title + meta.json ──────────


@pytest.mark.asyncio
@pytest.mark.unit
async def test_markdown_generation_minimal(tmp_path: Path) -> None:
    """LLM 直出 GFM markdown → 落盘 .md，验证 prompt 与 metadata。"""
    skill = SkillExecutor(_settings(tmp_path))
    md = (
        "# 英伟达 2025 Q3 分析\n\n"
        "## 执行摘要\n\n"
        "英伟达数据中心业务持续强势，毛利率维持高位。\n\n"
        "## 财务摘要\n\n"
        "| 指标 | Q2 | Q3 |\n"
        "| --- | --- | --- |\n"
        "| 营收 | 30B | 35B |\n"
        "| 毛利率 | 74% | 75% |\n\n"
        "## 风险\n\n"
        "- 客户集中度高\n- 供应链限制\n\n"
        "## 结论\n\n"
        "估值仍有空间。来源：英伟达 10-Q。\n" + "正文：" + "字" * 400
    )
    llm = FakeLLM(md)
    art = await skill.generate(llm=llm, artifact_type="markdown", brief="英伟达 2025 Q3 分析报告")
    assert art.artifact_type == "markdown"
    assert art.file_path.endswith(".md")
    assert Path(art.file_path).exists()
    assert Path(art.file_path).read_text(encoding="utf-8").startswith("#")
    assert int(art.metadata["heading_count"]) >= 3
    assert int(art.metadata["table_count"]) >= 1
    # title 来自 brief 前 40 字
    assert art.title.startswith("英伟达 2025 Q3")
    # prompt 是 markdown
    assert llm.last_messages is not None
    assert "Markdown" in llm.last_messages[0].content


@pytest.mark.asyncio
@pytest.mark.unit
async def test_markdown_alias_md_routes_same(tmp_path: Path) -> None:
    """artifact_type='md' 应归一为 markdown。"""
    skill = SkillExecutor(_settings(tmp_path))
    md = "# T\n\n## A\n\n## B\n\n" + "正文段落。" * 100
    llm = FakeLLM(md)
    art = await skill.generate(llm=llm, artifact_type="md", brief="t")
    assert art.artifact_type == "markdown"
    assert art.file_path.endswith(".md")


@pytest.mark.asyncio
@pytest.mark.unit
async def test_markdown_strip_outer_fence(tmp_path: Path) -> None:
    """LLM 把整篇 markdown 包在 ```markdown 围栏里时，应自动剥掉。"""
    skill = SkillExecutor(_settings(tmp_path))
    inner = "# Title\n\n## Sec\n\n" + "正文。" * 200
    md = "```markdown\n" + inner + "\n```"
    llm = FakeLLM(md)
    art = await skill.generate(llm=llm, artifact_type="markdown", brief="测试剥围栏")
    saved = Path(art.file_path).read_text(encoding="utf-8")
    assert saved.startswith("# Title")
    assert "```" not in saved


@pytest.mark.asyncio
@pytest.mark.unit
async def test_txt_generation_minimal(tmp_path: Path) -> None:
    """LLM 直出纯文本 → 落盘 .txt。"""
    skill = SkillExecutor(_settings(tmp_path))
    body = (
        "EchoDesk 每日待办 - 2026-05-28\n\n"
        "MORNING\n"
        "  - 跑 P4-M3 测试\n"
        "  - 同步 PR 进度\n\n"
        "AFTERNOON\n"
        "  - 整理英伟达分析草稿\n" + "  - 备注：上下文 ≥ 600 字符以通过健康检查；" * 20
    )
    llm = FakeLLM(body)
    art = await skill.generate(llm=llm, artifact_type="txt", brief="今天待办")
    assert art.artifact_type == "txt"
    assert art.file_path.endswith(".txt")
    saved = Path(art.file_path).read_text(encoding="utf-8")
    assert "EchoDesk" in saved
    assert int(art.metadata["line_count"]) >= 5


@pytest.mark.asyncio
@pytest.mark.unit
async def test_txt_too_short_raises(tmp_path: Path) -> None:
    skill = SkillExecutor(_settings(tmp_path))
    llm = FakeLLM("太短")
    with pytest.raises(SkillError, match="too short"):
        await skill.generate(llm=llm, artifact_type="txt", brief="x")


@pytest.mark.asyncio
@pytest.mark.unit
async def test_pdf_generation_minimal_ascii(tmp_path: Path) -> None:
    """PDF stub：用 fpdf2 内置 helvetica 不依赖中文字体，验证执行链路。

    不调 add_font，因为 CI 没字体；中文渲染由 test_pdf_generation_with_noto_font
    单独 @pytest.mark.requires_font 跳过 CI 没字体的情况。
    """
    pytest.importorskip("fpdf")
    skill = SkillExecutor(_settings(tmp_path))
    code = (
        "from fpdf import FPDF\n"
        "pdf = FPDF()\n"
        "pdf.add_page()\n"
        "pdf.set_font('helvetica', '', 12)\n"
        "pdf.cell(40, 10, 'Hello echo-demo')\n"
        "pdf.add_page()\n"
        "pdf.set_font('helvetica', '', 14)\n"
        "pdf.cell(40, 10, 'Page 2')\n"
        "pdf.output('output.pdf')\n"
    )
    llm = FakeLLM(code)
    art = await skill.generate(llm=llm, artifact_type="pdf", brief="生成最小 PDF")
    assert art.artifact_type == "pdf"
    assert art.file_path.endswith(".pdf")
    out = Path(art.file_path)
    assert out.exists()
    assert out.stat().st_size > 500
    # 文件头是 PDF 魔数
    assert out.read_bytes()[:4] == b"%PDF"
    # metadata
    assert int(art.metadata["pages_hint"]) >= 2
    # ECHODESK_PDF_FONT_PATH 应被注入子进程（哪怕代码没用）
    # ↑ 此处通过路由能跑成功间接确认；显式覆盖见
    #   test_pdf_executor_injects_font_env_var


@pytest.mark.asyncio
@pytest.mark.unit
async def test_pdf_executor_injects_font_env_var(tmp_path: Path) -> None:
    """PDF kind 走 exec_python_to_artifact 时，env 必须含 ECHODESK_PDF_FONT_PATH。"""
    pytest.importorskip("fpdf")
    skill = SkillExecutor(_settings(tmp_path))
    # 让 LLM 输出读取 env 变量并断言非空的代码 —— 跑成功证明环境变量被传进来了
    code = (
        "import os\n"
        "path = os.environ['ECHODESK_PDF_FONT_PATH']\n"
        "assert path and os.path.exists(path), f'font path missing: {path!r}'\n"
        "from fpdf import FPDF\n"
        "pdf = FPDF()\n"
        "pdf.add_page()\n"
        "pdf.set_font('helvetica', '', 12)\n"
        "pdf.cell(40, 10, 'env-injected')\n"
        "pdf.output('output.pdf')\n"
    )
    llm = FakeLLM(code)
    art = await skill.generate(llm=llm, artifact_type="pdf", brief="env injection 验证")
    assert Path(art.file_path).exists()


@pytest.mark.asyncio
@pytest.mark.unit
@pytest.mark.requires_font
async def test_pdf_generation_with_noto_font(tmp_path: Path) -> None:
    """中文 PDF：use add_font('noto', '', $ECHODESK_PDF_FONT_PATH) 渲染中文。

    需要 repo 内置 NotoSansSC-Regular.ttf；CI 没字体时通过
    ``@pytest.mark.requires_font`` 跳过（见 conftest）。
    """
    pytest.importorskip("fpdf")
    font_path = (
        Path(__file__).resolve().parents[2]
        / "app"
        / "adapters"
        / "skill"
        / "fonts"
        / "NotoSansSC-Regular.ttf"
    )
    if not font_path.exists():
        pytest.skip("NotoSansSC-Regular.ttf 未下载（开发环境 / CI）")
    skill = SkillExecutor(_settings(tmp_path))
    code = (
        "import os\n"
        "from fpdf import FPDF\n"
        "pdf = FPDF()\n"
        "pdf.add_page()\n"
        "pdf.add_font('noto', '', os.environ['ECHODESK_PDF_FONT_PATH'])\n"
        "pdf.set_font('noto', '', 14)\n"
        "pdf.cell(0, 10, '中文 PDF 验证：你好，世界！')\n"
        "pdf.output('output.pdf')\n"
    )
    llm = FakeLLM(code)
    art = await skill.generate(llm=llm, artifact_type="pdf", brief="中文 PDF")
    out = Path(art.file_path)
    assert out.exists()
    assert out.read_bytes()[:4] == b"%PDF"


@pytest.mark.asyncio
@pytest.mark.unit
async def test_artifact_meta_json_written(tmp_path: Path) -> None:
    """generate 完成后 build_dir/meta.json 必须含 title / artifact_type / ext。"""
    skill = SkillExecutor(_settings(tmp_path))
    md = "# T\n\n## S\n\n" + "正文段。" * 150
    llm = FakeLLM(md)
    art = await skill.generate(
        llm=llm, artifact_type="markdown", brief="英伟达 2025 财年第三季度业绩点评"
    )
    build_dir = Path(art.file_path).parent
    meta_path = build_dir / "meta.json"
    assert meta_path.exists()
    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    assert meta["title"].startswith("英伟达 2025 财年")
    assert meta["artifact_type"] == "markdown"
    assert meta["ext"] == "md"


@pytest.mark.asyncio
@pytest.mark.unit
async def test_html_meta_json_written(tmp_path: Path) -> None:
    """HTML 走的是 exec_text_to_file 同款路径，也要写 meta.json。"""
    skill = SkillExecutor(_settings(tmp_path))
    html = (
        "<!DOCTYPE html><html><head><script src='https://cdn.tailwindcss.com'></script>"
        "</head><body>" + "<p>正文段</p>" * 200 + "</body></html>"
    )
    llm = FakeLLM(html)
    art = await skill.generate(llm=llm, artifact_type="html", brief="生成 demo HTML 周报")
    meta_path = Path(art.file_path).parent / "meta.json"
    assert meta_path.exists()
    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    assert meta["artifact_type"] == "html"
    assert meta["ext"] == "html"
    assert "demo HTML" in meta["title"]


@pytest.mark.asyncio
@pytest.mark.unit
async def test_generated_artifact_has_non_empty_title(tmp_path: Path) -> None:
    """SkillExecutor 必须为每个产物填非空 title。"""
    skill = SkillExecutor(_settings(tmp_path))
    md = "# T\n\n## S\n\n" + "段。" * 150
    llm = FakeLLM(md)
    art = await skill.generate(llm=llm, artifact_type="markdown", brief="第一季度财报点评草稿")
    assert art.title == "第一季度财报点评草稿"
