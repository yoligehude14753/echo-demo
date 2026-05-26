"""Skill 4 格式真 LLM E2E（Yunwu M2.7）。

仅在 YUNWU_OPEN_KEY 配置且网络可达时跑，否则 skip。
pptx 需要 node/npm 在 PATH，缺则单独 skip。
"""

from __future__ import annotations

import os
import shutil
import socket
from pathlib import Path

import pytest
from app.adapters.llm.openai_compatible import OpenAICompatibleLLM
from app.adapters.skill.llm_skill import SkillExecutor
from app.config import Settings


def _yunwu_alive() -> bool:
    if not os.getenv("YUNWU_OPEN_KEY"):
        return False
    try:
        with socket.create_connection(("yunwu.ai", 443), timeout=3):
            return True
    except OSError:
        return False


pytestmark = pytest.mark.skipif(
    not _yunwu_alive(), reason="YUNWU_OPEN_KEY 未设置或网络不可达"
)


def _settings(tmp_path: Path) -> Settings:
    return Settings(
        storage_dir=tmp_path,
        skill_executor_build_dir=tmp_path / "skill_build",
        skill_executor_timeout_s=240,
        skill_executor_max_tokens=12_000,  # E2E 用小一点避免烧钱
    )


@pytest.mark.asyncio
@pytest.mark.integration
async def test_skill_html_real_yunwu(tmp_path: Path) -> None:
    s = _settings(tmp_path)
    llm = OpenAICompatibleLLM(s)
    skill = SkillExecutor(s)
    art = await skill.generate(
        llm=llm,
        artifact_type="html",
        brief=(
            "生成一份单文件 HTML 简报：英伟达 2020-2025 年营收快照（用占位真实数字），"
            "暗色主题 + Tailwind CDN + 至少一个 inline SVG 柱状图。无需联网。"
        ),
    )
    assert art.artifact_type == "html"
    assert art.file_path.endswith(".html")
    assert art.size_bytes > 1500
    html = Path(art.file_path).read_text(encoding="utf-8")
    assert "<!DOCTYPE" in html.upper() or "<html" in html.lower()
    assert "tailwind" in html.lower()
    await llm.aclose()


@pytest.mark.asyncio
@pytest.mark.integration
async def test_skill_word_real_yunwu(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    s = _settings(tmp_path)
    llm = OpenAICompatibleLLM(s)
    skill = SkillExecutor(s)
    art = await skill.generate(
        llm=llm,
        artifact_type="word",
        brief=(
            "生成一份 Word 投资简报：标题《英伟达 2025 投资展望》，"
            "包含执行摘要 + 2 个三级章节 + 至少一张 2×3 的表格 + 结论。全文中文 ≥ 800 字。"
        ),
    )
    assert art.artifact_type == "word"
    assert art.file_path.endswith(".docx")
    assert art.size_bytes > 5_000
    # docx 是 zip，能用 python-docx 打开
    from docx import Document

    doc = Document(art.file_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "英伟达" in text or "NVIDIA" in text.upper()
    assert len(doc.tables) >= 1
    await llm.aclose()


@pytest.mark.asyncio
@pytest.mark.integration
async def test_skill_xlsx_real_yunwu(tmp_path: Path) -> None:
    pytest.importorskip("openpyxl")
    s = _settings(tmp_path)
    llm = OpenAICompatibleLLM(s)
    skill = SkillExecutor(s)
    art = await skill.generate(
        llm=llm,
        artifact_type="xlsx",
        brief=(
            "生成 Excel 财务模型：英伟达 2020-2024 实际营收 + 2025-2026 预测，"
            "至少 3 个 sheet（假设 / 财务 / 预测），≥ 10 个公式单元格，"
            "≥ 2 个跨 sheet 引用。给出预测方法（增长率假设/DCF/或同比估算），"
            "公式必须能算（不能出现 #REF!）。"
        ),
    )
    assert art.artifact_type == "xlsx"
    assert art.file_path.endswith(".xlsx")
    assert art.size_bytes > 5_000
    from openpyxl import load_workbook

    wb = load_workbook(art.file_path, data_only=False)
    assert len(wb.sheetnames) >= 2
    # 找跨 sheet 公式
    cross_sheet = 0
    formula_cells = 0
    for sname in wb.sheetnames:
        ws = wb[sname]
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    formula_cells += 1
                    if "!" in cell.value:
                        cross_sheet += 1
    assert formula_cells >= 5
    await llm.aclose()


@pytest.mark.asyncio
@pytest.mark.integration
@pytest.mark.skipif(
    shutil.which("node") is None or shutil.which("npm") is None,
    reason="node / npm 未在 PATH（pptxgenjs 必需）",
)
async def test_skill_pptx_real_yunwu(tmp_path: Path) -> None:
    """真 Yunwu LLM 生成 pptxgenjs 代码 → 真跑 node → 验证 .pptx 文件。"""
    s = _settings(tmp_path)
    llm = OpenAICompatibleLLM(s)
    skill = SkillExecutor(s)
    art = await skill.generate(
        llm=llm,
        artifact_type="pptx",
        brief=(
            "生成 8 页 PPT：英伟达 2025 投资展望，16:9，封面 + 执行摘要 + "
            "3 个数据章节 + 至少一张数据表（如近 5 年营收对比）+ 结论。"
            "禁止 require('http')/require('fs')/require('child_process')。"
            "全程中文为主。"
        ),
    )
    assert art.artifact_type == "pptx"
    assert art.file_path.endswith(".pptx")
    assert art.size_bytes > 20_000  # pptx 是 zip，最小骨架约 25KB
    # pptx 是 zip 容器，简单校验
    import zipfile

    assert zipfile.is_zipfile(art.file_path)
    with zipfile.ZipFile(art.file_path) as zf:
        names = zf.namelist()
        assert "[Content_Types].xml" in names
        slide_files = [n for n in names if n.startswith("ppt/slides/slide")]
        assert len(slide_files) >= 5
    await llm.aclose()
