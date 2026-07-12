from __future__ import annotations

import argparse
import asyncio
import json
import multiprocessing
import os
import runpy
import sys
import traceback
import zipfile
from collections.abc import Sequence
from pathlib import Path
from typing import TextIO

# The frozen executable boundary must disable TorchScript before PyInstaller
# can load any collected application module.  ``app.__init__`` repeats this
# guard for source installs; keeping it here makes the packaged ordering an
# explicit invariant rather than relying on the current import graph.
os.environ["PYTORCH_JIT"] = "0"

PACKAGED_PYTHON_WORKER_FLAG = "--echodesk-python-worker"
ARTIFACT_RUNTIME_SMOKE_FLAG = "--artifact-runtime-smoke"

_SMOKE_FIXTURES: dict[str, tuple[str, str]] = {
    "docx": (
        "docx",
        "from docx import Document\n"
        "doc = Document()\n"
        "doc.add_heading('EchoDesk packaged runtime smoke', level=1)\n"
        "doc.add_paragraph('Deterministic DOCX artifact generated without an LLM.')\n"
        "doc.save('output.docx')\n",
    ),
    "xlsx": (
        "xlsx",
        "from openpyxl import Workbook\n"
        "wb = Workbook()\n"
        "ws = wb.active\n"
        "ws.title = 'Smoke'\n"
        "ws['A1'] = 'EchoDesk packaged runtime smoke'\n"
        "ws['B1'] = 301\n"
        "wb.save('output.xlsx')\n",
    ),
    "pdf": (
        "pdf",
        "from fpdf import FPDF\n"
        "pdf = FPDF()\n"
        "pdf.add_page()\n"
        "pdf.set_font('Helvetica', size=12)\n"
        "pdf.cell(0, 10, text='EchoDesk packaged runtime smoke')\n"
        "pdf.output('output.pdf')\n",
    ),
}


def _open_standard_stream(fd: int) -> TextIO:
    """Recreate redirected streams in PyInstaller's Windows noconsole mode."""

    try:
        return open(fd, "w", encoding="utf-8", buffering=1, closefd=False)
    except OSError:
        # The stream intentionally lives until this short-lived worker exits.
        return open(os.devnull, "w", encoding="utf-8")


def _ensure_worker_streams() -> None:
    # A noconsole PyInstaller executable sets these to None on Windows.  The
    # parent uses capture_output=True, so descriptors 1/2 are redirected pipes
    # and can be safely wrapped back into Python text streams.
    if sys.stdout is None:
        sys.stdout = _open_standard_stream(1)
    if sys.stderr is None:
        sys.stderr = _open_standard_stream(2)


def _write_stderr(message: str) -> None:
    _ensure_worker_streams()
    assert sys.stderr is not None
    sys.stderr.write(message)
    sys.stderr.flush()


def _write_stdout(message: str) -> None:
    _ensure_worker_streams()
    assert sys.stdout is not None
    sys.stdout.write(message)
    sys.stdout.flush()


def _ensure_source_backend_path() -> None:
    if getattr(sys, "frozen", False):
        return
    backend_root = str(Path(__file__).resolve().parent.parent)
    if backend_root not in sys.path:
        sys.path.insert(0, backend_root)


def _resolve_worker_script(raw_path: str) -> Path:
    candidate = Path(raw_path).expanduser()
    if not candidate.is_absolute():
        raise ValueError("worker script path must be absolute")
    if candidate.is_symlink():
        raise ValueError("worker script must not be a symbolic link")
    try:
        script_path = candidate.resolve(strict=True)
    except OSError as exc:
        raise ValueError(f"worker script is unavailable: {exc}") from exc
    if not script_path.is_file() or script_path.suffix.lower() != ".py":
        raise ValueError("worker script must be a regular .py file")
    if script_path.parent != Path.cwd().resolve():
        raise ValueError("worker script must be located directly in the worker cwd")
    return script_path


def _system_exit_code(value: object) -> int:
    if value is None:
        return 0
    if isinstance(value, int):
        return value
    _write_stderr(f"{value}\n")
    return 1


def run_python_worker(script: str, script_args: Sequence[str] = ()) -> int:
    """Execute one generated script with normal Python CLI semantics."""

    try:
        script_path = _resolve_worker_script(script)
    except ValueError as exc:
        _write_stderr(f"EchoDesk artifact worker refused script: {exc}\n")
        return 2

    _ensure_worker_streams()
    os.chdir(script_path.parent)
    sys.argv = [str(script_path), *script_args]
    script_directory = str(script_path.parent)
    if not sys.path or sys.path[0] != script_directory:
        sys.path.insert(0, script_directory)
    try:
        runpy.run_path(str(script_path), run_name="__main__")
    except SystemExit as exc:
        return _system_exit_code(exc.code)
    except KeyboardInterrupt:
        _write_stderr("KeyboardInterrupt\n")
        return 130
    except BaseException:
        _write_stderr(traceback.format_exc())
        return 1
    return 0


async def _generate_smoke_artifacts(output_dir: Path) -> dict[str, Path]:
    from app.adapters.skill.python_executor import exec_python_to_artifact

    artifacts: dict[str, Path] = {}
    for kind, (extension, code) in _SMOKE_FIXTURES.items():
        result = await exec_python_to_artifact(
            code,
            output_dir / kind,
            expected_ext=extension,
            timeout_s=60.0,
        )
        if not result.success or result.output_path is None:
            raise RuntimeError(f"{kind} packaged runtime failed: {result.stderr}")
        artifacts[kind] = result.output_path
    return artifacts


def _validate_smoke_artifacts(artifacts: dict[str, Path]) -> dict[str, int]:
    import pdfplumber
    from docx import Document
    from openpyxl import load_workbook  # type: ignore[import-untyped]

    doc = Document(str(artifacts["docx"]))
    if not any("EchoDesk packaged runtime smoke" in p.text for p in doc.paragraphs):
        raise RuntimeError("DOCX packaged runtime smoke content is unreadable")

    workbook = load_workbook(artifacts["xlsx"], read_only=True, data_only=False)
    try:
        if workbook["Smoke"]["A1"].value != "EchoDesk packaged runtime smoke":
            raise RuntimeError("XLSX packaged runtime smoke content is unreadable")
    finally:
        workbook.close()

    with pdfplumber.open(artifacts["pdf"]) as pdf:
        if len(pdf.pages) != 1 or "EchoDesk packaged runtime smoke" not in (
            pdf.pages[0].extract_text() or ""
        ):
            raise RuntimeError("PDF packaged runtime smoke content is unreadable")

    return {kind: path.stat().st_size for kind, path in artifacts.items()}


def _generate_ppt_smoke_artifact(output_dir: Path) -> Path:
    from app.adapters.skill.llm_skill import run_packaged_ppt_runtime_smoke
    from app.config import Settings

    settings = Settings(_env_file=None)  # type: ignore[call-arg]
    return run_packaged_ppt_runtime_smoke(
        output_dir / "pptx",
        node_bin=settings.resolved_skill_node_bin,
        electron_runtime=settings.resolved_skill_node_is_electron,
        timeout_s=60.0,
    )


def _write_epub_smoke(path: Path) -> None:
    marker = "EchoDesk parser runtime smoke"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "mimetype",
            "application/epub+zip",
            compress_type=zipfile.ZIP_STORED,
        )
        archive.writestr(
            "META-INF/container.xml",
            "<?xml version='1.0'?>"
            "<container version='1.0' xmlns='urn:oasis:names:tc:opendocument:xmlns:container'>"
            "<rootfiles><rootfile full-path='OEBPS/content.opf' "
            "media-type='application/oebps-package+xml'/></rootfiles></container>",
        )
        archive.writestr(
            "OEBPS/content.opf",
            "<?xml version='1.0'?>"
            "<package version='2.0' xmlns='http://www.idpf.org/2007/opf' unique-identifier='id'>"
            "<metadata xmlns:dc='http://purl.org/dc/elements/1.1/'>"
            "<dc:title>EchoDesk Smoke</dc:title><dc:identifier id='id'>smoke</dc:identifier>"
            "<dc:language>en</dc:language></metadata>"
            "<manifest><item id='chapter' href='chapter.xhtml' "
            "media-type='application/xhtml+xml'/></manifest>"
            "<spine><itemref idref='chapter'/></spine></package>",
        )
        archive.writestr(
            "OEBPS/chapter.xhtml",
            f"<html xmlns='http://www.w3.org/1999/xhtml'><body><p>{marker}</p></body></html>",
        )


def _validate_rag_parser_runtime(
    output_dir: Path,
    artifacts: dict[str, Path],
) -> dict[str, int]:
    from app.adapters.diarizer.ecapa import ECAPADiarizer
    from app.adapters.rag.parsers import parse_to_text

    del ECAPADiarizer  # import is the frozen ECAPA dependency boundary
    marker = "EchoDesk parser runtime smoke"
    html_path = output_dir / "parser-smoke.html"
    csv_path = output_dir / "parser-smoke.csv"
    epub_path = output_dir / "parser-smoke.epub"
    html_path.write_text(f"<html><body><p>{marker}</p></body></html>", encoding="utf-8")
    csv_path.write_text(f"kind,value\nparser,{marker}\n", encoding="utf-8")
    _write_epub_smoke(epub_path)

    parser_inputs = {
        "docx": artifacts["docx"],
        "xlsx": artifacts["xlsx"],
        "pdf": artifacts["pdf"],
        "pptx": artifacts["pptx"],
        "html": html_path,
        "csv": csv_path,
        "epub": epub_path,
    }
    parsed_chars: dict[str, int] = {}
    for kind, path in parser_inputs.items():
        text = parse_to_text(path)
        if not text.strip():
            raise RuntimeError(f"{kind} packaged RAG parser returned empty content")
        if kind in {"html", "csv", "epub"} and marker not in text:
            raise RuntimeError(f"{kind} packaged RAG parser lost its marker")
        parsed_chars[kind] = len(text)
    return parsed_chars


def _validate_cpu_diarizer_runtime() -> dict[str, object]:
    """Exercise the frozen ECAPA dependency boundary without fetching a model."""

    import torch
    import torch.distributed as distributed
    import torchaudio
    from speechbrain.inference.speaker import SpeakerRecognition

    del SpeakerRecognition  # the model class import is the frozen SpeechBrain boundary
    cuda_build = torch.version.cuda
    cuda_available = bool(torch.cuda.is_available())
    jit_enabled = bool(torch.jit._state._enabled)
    if cuda_build is not None or cuda_available:
        raise RuntimeError(
            "packaged diarizer must use the official CPU torch runtime "
            f"(cuda_build={cuda_build!r}, cuda_available={cuda_available})"
        )
    if jit_enabled:
        raise RuntimeError("packaged diarizer unexpectedly enabled TorchScript")

    with torch.no_grad():
        probe = torch.tensor([[3.0, 4.0]], dtype=torch.float32)
        normalized = torch.nn.functional.normalize(probe, dim=1)
        norm = float(torch.linalg.vector_norm(normalized).item())
    if abs(norm - 1.0) > 1e-6:
        raise RuntimeError(f"packaged CPU eager tensor probe returned invalid norm {norm}")

    return {
        "cpu_only": True,
        "cuda_available": cuda_available,
        "cuda_build": cuda_build,
        "distributed_available": bool(distributed.is_available()),
        "jit_enabled": jit_enabled,
        "torch_version": str(torch.__version__),
        "torchaudio_version": str(torchaudio.__version__),
        "vector_norm": norm,
    }


def run_artifact_runtime_smoke(raw_output_dir: str) -> int:
    """Generate and reopen deterministic artifacts using the packaged runtime."""

    output_dir = Path(raw_output_dir).expanduser().resolve()
    try:
        _ensure_source_backend_path()
        output_dir.mkdir(parents=True, exist_ok=True)
        artifacts = asyncio.run(_generate_smoke_artifacts(output_dir))
        artifacts["pptx"] = _generate_ppt_smoke_artifact(output_dir)
        sizes = _validate_smoke_artifacts(artifacts)
        parser_chars = _validate_rag_parser_runtime(output_dir, artifacts)
        diarizer_runtime = _validate_cpu_diarizer_runtime()
        manifest = {
            "ok": True,
            "output_dir": str(output_dir),
            "artifacts": {
                kind: {"path": str(artifacts[kind]), "size_bytes": sizes[kind]}
                for kind in sorted(artifacts)
            },
            "diarizer_runtime": diarizer_runtime,
            "rag_parser_chars": parser_chars,
        }
        (output_dir / "artifact-runtime-smoke.json").write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        _write_stdout(json.dumps(manifest, ensure_ascii=False) + "\n")
    except Exception:
        _write_stderr(traceback.format_exc())
        return 1
    return 0


def _server_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="EchoDesk bundled backend")
    parser.add_argument("--host", default=os.getenv("ECHO_BACKEND_BIND_HOST", "127.0.0.1"))
    parser.add_argument(
        "--port",
        type=int,
        default=int(os.getenv("ECHO_BACKEND_PORT", "8769")),
    )
    parser.add_argument("--log-level", default="info")
    parser.add_argument("--ws-max-size", type=int, choices=(4096,), default=4096)
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    args = list(sys.argv[1:] if argv is None else argv)
    if args and args[0] == PACKAGED_PYTHON_WORKER_FLAG:
        worker_parser = argparse.ArgumentParser(add_help=False)
        worker_parser.add_argument("script")
        worker_parser.add_argument("script_args", nargs=argparse.REMAINDER)
        worker_args = worker_parser.parse_args(args[1:])
        return run_python_worker(worker_args.script, worker_args.script_args)
    if args and args[0] == ARTIFACT_RUNTIME_SMOKE_FLAG:
        smoke_parser = argparse.ArgumentParser(add_help=False)
        smoke_parser.add_argument("output_dir")
        smoke_args = smoke_parser.parse_args(args[1:])
        return run_artifact_runtime_smoke(smoke_args.output_dir)

    server_args = _server_parser().parse_args(args)
    _ensure_source_backend_path()
    import uvicorn

    uvicorn.run(
        "app.main:app",
        host=server_args.host,
        port=server_args.port,
        log_level=server_args.log_level,
        access_log=False,
        factory=False,
        ws_max_size=server_args.ws_max_size,
    )
    return 0


if __name__ == "__main__":
    multiprocessing.freeze_support()
    raise SystemExit(main())
