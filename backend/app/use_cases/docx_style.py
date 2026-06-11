"""use_case: 参考 docx 样式抽离与模仿（Phase 2 of 政务/制式文档 skill）。

用户原话（2026-06-05）："风格模仿、格式模仿、字体等内容的模仿和抽离"。

做法：读一份参考 .docx，抽出可复刻的版式规格（StyleSpec）——页边距、正文字体/
字号/行距、各级标题字体字号、是否红头、主题色——再把它转成中文指令注入到文档
生成提示词（走现有 ``ArtifactRequest.extra_instructions``）。这样生成的新文档能
复刻参考件的"形式"，而内容仍由 brief 决定。

只依赖 python-docx；纯读取，无副作用，可离线单测。
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger("echodesk.docx_style")

# 抽离时只看正文前 N 段，足够覆盖标题/正文/各级小标题，避免长文档全扫。
_SCAN_PARAGRAPHS = 80


@dataclass(slots=True)
class FontSpec:
    """单处文字的字体规格（eastAsia 中文字体优先，西文兜底）。"""

    east_asia: str | None = None
    latin: str | None = None
    size_pt: float | None = None
    bold: bool = False
    color_hex: str | None = None  # RRGGBB，无则 None

    def label(self) -> str:
        name = self.east_asia or self.latin or "默认"
        parts = [name]
        if self.size_pt:
            parts.append(f"{self.size_pt:g}pt")
        if self.bold:
            parts.append("加粗")
        if self.color_hex and self.color_hex.lower() not in ("000000", "auto"):
            parts.append(f"#{self.color_hex}")
        return " ".join(parts)


@dataclass(slots=True)
class StyleSpec:
    """从参考 docx 抽出的可复刻版式规格。"""

    margins_mm: dict[str, float] = field(default_factory=dict)  # top/bottom/left/right
    page_size_mm: dict[str, float] = field(default_factory=dict)  # width/height
    body: FontSpec = field(default_factory=FontSpec)
    title: FontSpec | None = None  # 居中大标题（含红色判定）
    headings: dict[int, FontSpec] = field(default_factory=dict)  # level -> font
    line_spacing_pt: float | None = None  # 正文固定行距（磅）
    has_red_head: bool = False  # 是否含红色发文机关标志（红头）

    def to_instructions(self) -> str:
        """转成注入生成提示词的中文样式指令（严格复刻参考件的形式）。"""
        lines: list[str] = ["【严格复刻参考文档样式（仅形式，内容仍按 brief）】"]
        m = self.margins_mm
        if m:
            lines.append(
                f"- 页边距(mm)：上{m.get('top', 0):.0f} 下{m.get('bottom', 0):.0f} "
                f"左{m.get('left', 0):.0f} 右{m.get('right', 0):.0f}"
            )
        if self.title is not None:
            lines.append(f"- 大标题：{self.title.label()}，居中")
        if self.body.east_asia or self.body.size_pt:
            extra = f"，行距固定{self.line_spacing_pt:g}磅" if self.line_spacing_pt else ""
            lines.append(f"- 正文：{self.body.label()}{extra}")
        for level in sorted(self.headings):
            lines.append(f"- {level} 级标题：{self.headings[level].label()}")
        if self.has_red_head:
            lines.append("- 保留红头（发文机关标志红色、红色分隔线）")
        lines.append("- 字体/字号/边距必须与上面一致；缺失项沿用合理默认，不要自创风格。")
        return "\n".join(lines)


def _run_font(run: object) -> FontSpec:
    rpr = getattr(run, "_element", None)
    spec = FontSpec()
    try:
        spec.size_pt = run.font.size.pt if run.font.size else None  # type: ignore[attr-defined]
        spec.bold = bool(run.font.bold)  # type: ignore[attr-defined]
        spec.latin = run.font.name  # type: ignore[attr-defined]
        color = run.font.color  # type: ignore[attr-defined]
        if color is not None and color.rgb is not None:
            spec.color_hex = str(color.rgb)
    except Exception:
        pass
    try:
        if rpr is not None and rpr.rPr is not None and rpr.rPr.rFonts is not None:
            ea = rpr.rPr.rFonts.get(qn("w:eastAsia"))
            if ea:
                spec.east_asia = ea
    except Exception:
        pass
    return spec


def _is_reddish(color_hex: str | None) -> bool:
    if not color_hex or len(color_hex) != 6:
        return False
    try:
        r = int(color_hex[0:2], 16)
        g = int(color_hex[2:4], 16)
        b = int(color_hex[4:6], 16)
    except ValueError:
        return False
    return r >= 0xA0 and g <= 0x60 and b <= 0x60


def extract_docx_style(path: str) -> StyleSpec:
    """读参考 .docx，抽出可复刻的 StyleSpec。失败时返回尽量填好的部分结果。"""
    doc = Document(path)
    spec = StyleSpec()

    section = doc.sections[0]
    try:
        spec.margins_mm = {
            "top": float(section.top_margin.mm),
            "bottom": float(section.bottom_margin.mm),
            "left": float(section.left_margin.mm),
            "right": float(section.right_margin.mm),
        }
        spec.page_size_mm = {
            "width": float(section.page_width.mm),
            "height": float(section.page_height.mm),
        }
    except Exception:
        pass

    body_candidates: list[FontSpec] = []
    for p in doc.paragraphs[:_SCAN_PARAGRAPHS]:
        text = p.text.strip()
        if not text:
            continue
        runs = [r for r in p.runs if r.text.strip()]
        if not runs:
            continue
        font = _run_font(runs[0])
        style_name = (p.style.name or "").lower() if p.style else ""
        is_centered = p.alignment is not None and int(p.alignment) == 1  # CENTER

        # 标题：居中 + 字号偏大（或样式名 title），取首个命中
        if spec.title is None and is_centered and (font.size_pt or 0) >= 18:
            spec.title = font
            if _is_reddish(font.color_hex):
                spec.has_red_head = True
            continue

        # 各级标题：python-docx 内置 Heading N 样式
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
            spec.headings.setdefault(level, font)
            continue

        # 正文候选
        body_candidates.append(font)
        if p.paragraph_format.line_spacing is not None and spec.line_spacing_pt is None:
            try:
                ls = p.paragraph_format.line_spacing
                spec.line_spacing_pt = float(ls.pt) if hasattr(ls, "pt") else None
            except Exception:
                pass

    if body_candidates:
        # 取出现最多的正文字体作为代表
        spec.body = max(
            body_candidates,
            key=lambda f: sum(
                1 for c in body_candidates if c.east_asia == f.east_asia and c.size_pt == f.size_pt
            ),
        )

    # 红头兜底：标题非红，但前几段有红色大字也算
    if not spec.has_red_head:
        spec.has_red_head = _detect_red_head(doc)

    return spec


def _detect_red_head(doc: object) -> bool:
    """前 12 段里是否出现红色大字（≥16pt）—— 红头发文机关标志的特征。"""
    for p in doc.paragraphs[:12]:  # type: ignore[attr-defined]
        for r in p.runs:
            f = _run_font(r)
            if _is_reddish(f.color_hex) and (f.size_pt or 0) >= 16:
                return True
    return False


def extract_style_instructions(path: str) -> str:
    """便捷入口：抽样式 → 直接返回可注入 ``extra_instructions`` 的中文指令。"""
    return extract_docx_style(path).to_instructions()
